from array import *
import sys
import pandas as pd
import math
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.pyplot import figure
from scipy.optimize import curve_fit
import seaborn as sns

a1 = sys.argv[1]

data = pd.read_csv(a1)

y = data.iloc[:,0].values

#finding median
y_meadian = np.median(y)

# First quartile (Q1) 
Q1 = np.percentile(y, 25, interpolation = 'midpoint') 
#print(Q1)
  
# Third quartile (Q3) 
Q3 = np.percentile(y, 75, interpolation = 'midpoint') 
#print(Q3)
  
# Interquaritle range (IQR) 
IQR = Q3 - Q1 
#print(IQR)

#A commonly used rule says that a data point is an outlier if it is more than Q3 + (IQR*1.5)
# or less than Q1-(IQR*1.5)
upper_limit = Q3 + (IQR*1.5)
#print(upper_limit)
lower_limit = Q1 - (IQR*1.5)
#print(lower_limit)

outliers = []

for i in range(len(y)):
    if(y[i]>upper_limit or y[i]< lower_limit ):
        outliers.append(y[i])
        
print(outliers)

figure(num=None, figsize=(5, 6), dpi=80, facecolor='w', edgecolor='k')
sns.boxplot(y)
plt.title('boxplot')
plt.savefig('boxplot_for_outliers.png')


from docx import Document
from docx.shared import Inches

document = Document()

p = document.add_paragraph()
r= p.add_run()
r.add_picture('boxplot_for_outliers.png')

p = document.add_paragraph()
r= p.add_run()
r.add_text("The outliers of the given data are: ")

table = document.add_table(rows=1, cols=1)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'outlier'

for outlier in outliers:
    row_cells = table.add_row().cells
    row_cells[0].text = str(outlier)



document.save('Outliers.docx')








